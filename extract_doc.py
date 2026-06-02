#!/usr/bin/env python
# -*- coding: utf-8 -*-
import sys

try:
    from docx import Document
    
    doc = Document('Especificação do Sistema - Programação Avançada (1).docx')
    
    print("=" * 80)
    print("CONTEÚDO DO DOCUMENTO")
    print("=" * 80)
    
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    print("\n" + "=" * 80)
    print("TABELAS ENCONTRADAS")
    print("=" * 80)
    
    for i, table in enumerate(doc.tables):
        print(f"\n--- TABELA {i+1} ---")
        for row in table.rows:
            print(" | ".join(cell.text for cell in row.cells))
            
except ImportError:
    print("python-docx não está instalado. Instalando...")
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    print("Tente novamente.")
except Exception as e:
    print(f"Erro: {e}")
    import traceback
    traceback.print_exc()

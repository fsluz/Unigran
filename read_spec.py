#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Extrai e processa o arquivo DOCX anexado
"""
import sys
import os
from docx import Document

# Tenta vários caminhos possíveis
caminhos_possiveis = [
    "Especificação do Sistema - Programação Avançada (1).docx",
    "Especificação do Sistema - Programação Avançada.docx",
    "especificacao.docx",
]

# Procura em temp também
import tempfile
temp_dir = tempfile.gettempdir()

doc_encontrado = None

for caminho in caminhos_possiveis:
    if os.path.exists(caminho):
        doc_encontrado = caminho
        break
    elif os.path.exists(os.path.join(temp_dir, caminho)):
        doc_encontrado = os.path.join(temp_dir, caminho)
        break

if not doc_encontrado:
    # Listar TODOS os .docx
    print("Arquivos .docx encontrados no sistema:")
    for root, dirs, files in os.walk("."):
        for file in files:
            if file.endswith(".docx"):
                print(f"  {os.path.join(root, file)}")
    sys.exit(1)

try:
    doc = Document(doc_encontrado)
    print(f"Especificação carregada: {doc_encontrado}\n")
    print("="*80)
    print("CONTEÚDO COMPLETO")
    print("="*80)
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            nivel = len(para.style.name) - len(para.style.name.lstrip('Heading'))
            prefix = "  " * max(0, nivel - 1)
            print(f"{prefix}{para.text}")
    
    print("\n" + "="*80)
    print("TABELAS")
    print("="*80)
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTABELA {table_idx + 1}:")
        for row in table.rows:
            print(" | ".join(cell.text.strip() for cell in row.cells))
            
except Exception as e:
    print(f"Erro: {e}")
    import traceback
    traceback.print_exc()
